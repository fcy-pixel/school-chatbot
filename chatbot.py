"""
School Affairs Chatbot — Core Logic (RAG)
Strategy:
  1. On startup: download ALL Word docs, split into large overlapping chunks,
      build an in-memory keyword index (full-text, not just summaries).
  2. Per query: score every chunk by CJK character + Latin word overlap;
      pick the top-K most relevant chunks from across all documents.
  3. Send only those chunks to Qwen — accurate and efficient.
"""

import io
import re
import base64
import requests
from typing import Optional
from openai import OpenAI
from docx import Document as DocxDocument

# ── Constants ──────────────────────────────────────────────────────────────────
QWEN_BASE_URL = "https://dashscope-intl.aliyuncs.com/compatible-mode/v1"
DEFAULT_MODEL  = "qwen-plus"
CHUNK_SIZE     = 3000   # large chunks = more context per hit
CHUNK_OVERLAP  = 500    # overlap ensures topics at chunk boundaries are captured
TOP_K          = 8      # chunks sent to the model per query
MAX_CTX_CHARS  = 24_000 # hard cap on total context chars sent to model


class ChatbotError(Exception):
    pass


class Chunk:
    """A text chunk from a document with pre-tokenised keyword sets for fast scoring."""
    __slots__ = ("text", "source", "page_tag", "tok_cjk", "tok_latin")

    def __init__(self, text: str, source: str, page_tag: str):
        self.text      = text
        self.source    = source
        self.page_tag  = page_tag
        self.tok_cjk   = {c for c in text if "\u4e00" <= c <= "\u9fff"}
        self.tok_latin = set(re.findall(r"[a-zA-Z0-9]+", text.lower()))


class SchoolChatbot:
    def __init__(
        self,
        qwen_api_key: str,
        github_repo: str,
        github_path: str = "",
        github_token: Optional[str] = None,
        model: str = DEFAULT_MODEL,
    ):
        if not qwen_api_key:
            raise ChatbotError("缺少 Qwen API Key")

        self.client       = OpenAI(api_key=qwen_api_key, base_url=QWEN_BASE_URL)
        self.github_repo  = github_repo.strip("/") if github_repo else ""
        self.github_path  = github_path.strip("/") if github_path else ""
        self.github_token = github_token
        self.model        = model

        self._doc_text_cache: dict[str, str]       = {}
        self._doc_list_cache: Optional[list[dict]] = None
        self._uploaded_docs:  set[str]             = set()
        self._chunk_index:    list[Chunk]           = []
        self._indexed_docs:   set[str]             = set()

    # ── Properties ─────────────────────────────────────────────────────────────

    @property
    def has_content(self) -> bool:
        return bool(self._doc_list_cache or self._uploaded_docs)

    @property
    def index_ready(self) -> bool:
        return bool(self._chunk_index)

    # ── GitHub helpers ─────────────────────────────────────────────────────────

    def _gh_headers(self) -> dict:
        h = {"Accept": "application/vnd.github.v3+json"}
        if self.github_token:
            h["Authorization"] = f"token {self.github_token}"
        return h

    def get_doc_list(self, force_refresh: bool = False) -> list[dict]:
        if self._doc_list_cache is not None and not force_refresh:
            return self._doc_list_cache

        path = self.github_path or ""
        url  = f"https://api.github.com/repos/{self.github_repo}/contents/{path}"

        try:
            resp = requests.get(url, headers=self._gh_headers(), timeout=15)
        except requests.RequestException as e:
            raise ChatbotError(f"GitHub 連接失敗：{e}")

        if resp.status_code == 404:
            raise ChatbotError(f"找不到倉庫或路徑：{self.github_repo}/{path}")
        if resp.status_code == 403:
            raise ChatbotError("GitHub API 速率限制，請稍後再試或提供 Personal Access Token")
        if not resp.ok:
            raise ChatbotError(f"GitHub API 錯誤 {resp.status_code}")

        items = resp.json()
        if not isinstance(items, list):
            raise ChatbotError("指定路徑不是目錄，請確認文件夾路徑")

        self._doc_list_cache = [
            {
                "name":         item["name"],
                "download_url": item["download_url"],
                "path":         item["path"],
                "size":         item.get("size", 0),
            }
            for item in items
            if item.get("type") == "file" and item["name"].lower().endswith(".docx")
        ]
        return self._doc_list_cache

    def push_doc_to_github(self, filename: str, data: bytes) -> str:
        if not self.github_repo:
            raise ChatbotError("未設定 GitHub 倉庫，無法儲存文件")
        if not self.github_token:
            raise ChatbotError("需要 GitHub Token（具備 repo 寫入權限）才能儲存文件")

        file_path = f"{self.github_path}/{filename}" if self.github_path else filename
        api_url   = f"https://api.github.com/repos/{self.github_repo}/contents/{file_path}"
        headers   = {**self._gh_headers(), "Content-Type": "application/json"}

        sha: Optional[str] = None
        check = requests.get(api_url, headers=headers, timeout=10)
        if check.status_code == 200:
            sha = check.json().get("sha")

        payload: dict = {
            "message": f"上傳文件：{filename}",
            "content": base64.b64encode(data).decode(),
        }
        if sha:
            payload["sha"] = sha

        resp = requests.put(api_url, headers=headers, json=payload, timeout=30)
        if resp.status_code in (200, 201):
            self._doc_list_cache = None
            return resp.json().get("content", {}).get("html_url", "")
        elif resp.status_code == 401:
            raise ChatbotError("GitHub Token 無效或已過期，請重新填寫")
        elif resp.status_code == 403:
            raise ChatbotError("GitHub Token 沒有倉庫寫入權限，需要 repo scope")
        else:
            raise ChatbotError(
                f"上傳失敗（{resp.status_code}）：{resp.json().get('message', resp.text[:200])}"
            )

    # ── Delete document ──────────────────────────────────────────────────────

    def delete_doc(self, filename: str) -> None:
        """Remove a document from the in-memory index and from GitHub (if configured)."""
        # Remove from chunk index and all caches
        self._chunk_index = [c for c in self._chunk_index if c.source != filename]
        self._doc_text_cache.pop(filename, None)
        self._uploaded_docs.discard(filename)
        self._indexed_docs.discard(filename)
        if self._doc_list_cache is not None:
            self._doc_list_cache = [d for d in self._doc_list_cache if d["name"] != filename]

        # Delete from GitHub if configured
        if not self.github_repo or not self.github_token:
            return
        file_path = f"{self.github_path}/{filename}" if self.github_path else filename
        api_url   = f"https://api.github.com/repos/{self.github_repo}/contents/{file_path}"
        headers   = {**self._gh_headers(), "Content-Type": "application/json"}

        check = requests.get(api_url, headers=headers, timeout=10)
        if check.status_code == 404:
            return  # already gone
        if check.status_code != 200:
            raise ChatbotError(f"查詢 GitHub 文件失敗（{check.status_code}）")

        sha  = check.json().get("sha")
        resp = requests.delete(
            api_url,
            headers=headers,
            json={"message": f"刪除文件：{filename}", "sha": sha},
            timeout=15,
        )
        if not resp.ok:
            raise ChatbotError(
                f"GitHub 刪除失敗（{resp.status_code}）：{resp.json().get('message', '')}"
            )

    # ── Word extraction ──────────────────────────────────────────────────────

    def _extract_text(self, doc: dict) -> str:
        """Download a .docx file and return its full extracted text. Result is cached."""
        name = doc["name"]
        if name in self._doc_text_cache:
            return self._doc_text_cache[name]

        try:
            resp = requests.get(doc["download_url"], timeout=30)
            resp.raise_for_status()
            document = DocxDocument(io.BytesIO(resp.content))
            paragraphs = []
            for para in document.paragraphs:
                t = para.text.strip()
                if t:
                    paragraphs.append(t)
            # Also extract text from tables
            for table in document.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            result = "\n".join(paragraphs)
            if not result.strip():
                result = f"[{name} 無法提取文字]"
        except Exception as e:
            result = f"[無法讀取 {name}：{e}]"

        self._doc_text_cache[name] = result
        return result

    # ── Chunk index ────────────────────────────────────────────────────────────

    def _make_chunks(self, text: str, source: str) -> list[Chunk]:
        page_re      = re.compile(r"\[第\s*\d+\s*頁\]")
        current_page = ""
        chunks: list[Chunk] = []
        start = 0
        while start < len(text):
            end   = min(start + CHUNK_SIZE, len(text))
            piece = text[start:end]
            for m in page_re.finditer(text, 0, end):
                current_page = m.group()
            chunks.append(Chunk(piece, source, current_page))
            if end >= len(text):
                break
            start += CHUNK_SIZE - CHUNK_OVERLAP
        return chunks

    def build_index(self, progress_callback=None) -> int:
        """
        Download every Word doc from GitHub, extract full text, build chunk index.
        Call once on startup; queries will be fast thereafter.
        progress_callback(done, total, filename) — optional UI progress hook.
        """
        self._chunk_index.clear()
        self._indexed_docs.clear()

        docs  = self.get_doc_list(force_refresh=True)
        total = len(docs)

        for i, doc in enumerate(docs):
            if progress_callback:
                progress_callback(i, total, doc["name"])
            text = self._extract_text(doc)
            if not text.startswith("[無法") and "無法提取文字" not in text:
                self._chunk_index.extend(self._make_chunks(text, doc["name"]))
                self._indexed_docs.add(doc["name"])

        if progress_callback:
            progress_callback(total, total, "完成")
        return len(self._chunk_index)

    def _score_chunk(self, chunk: Chunk, q_cjk: set, q_latin: set) -> float:
        cjk = len(q_cjk & chunk.tok_cjk)    / len(q_cjk)   if q_cjk   else 0.0
        lat = len(q_latin & chunk.tok_latin) / len(q_latin) if q_latin else 0.0
        return max(cjk, lat)

    def search_index(self, query: str) -> list[Chunk]:
        """Return the top-K most relevant chunks for the given query."""
        q_cjk   = {c for c in query if "\u4e00" <= c <= "\u9fff"}
        q_latin = set(re.findall(r"[a-zA-Z0-9]+", query.lower()))
        scored  = sorted(
            ((c, self._score_chunk(c, q_cjk, q_latin)) for c in self._chunk_index),
            key=lambda x: x[1],
            reverse=True,
        )
        results: list[Chunk] = []
        total_chars = 0
        for chunk, score in scored:
            if score <= 0 or total_chars >= MAX_CTX_CHARS:
                break
            results.append(chunk)
            total_chars += len(chunk.text)
            if len(results) >= TOP_K:
                break
        return results

    # ── Word upload ────────────────────────────────────────────────────────────

    def ingest_uploaded_doc(self, name: str, data: bytes) -> int:
        """Extract text from an uploaded .docx file and add to the chunk index."""
        try:
            document   = DocxDocument(io.BytesIO(data))
            paragraphs = []
            for para in document.paragraphs:
                t = para.text.strip()
                if t:
                    paragraphs.append(t)
            for table in document.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            text = "\n".join(paragraphs)
            if not text.strip():
                text = f"[{name} 無法提取文字]"
        except Exception as e:
            text = f"[無法讀取 {name}：{e}]"

        self._doc_text_cache[name] = text
        self._uploaded_docs.add(name)

        # Remove stale chunks for this file then add fresh ones
        self._chunk_index = [c for c in self._chunk_index if c.source != name]
        if not text.startswith("[無法") and "無法提取文字" not in text:
            new_chunks = self._make_chunks(text, name)
            self._chunk_index.extend(new_chunks)
            return len(new_chunks)
        return 0

    # ── Public chat API ────────────────────────────────────────────────────────

    def chat(
        self,
        question: str,
        chat_history: Optional[list[dict]] = None,
    ) -> tuple[str, list[str]]:
        """
        Answer a question using top-K chunks from the full-text index.
        """
        doc_list = self.get_doc_list() if self.github_repo else []
        if not doc_list and not self._uploaded_docs:
            return "目前沒有任何文件可供查閱，請管理員上傳 Word 文件或設定 GitHub 倉庫。", []

        context_parts: list[str] = []
        source_files:  list[str] = []

        if self.index_ready:
            top_chunks = self.search_index(question)
            if not top_chunks:
                context_parts = ["（索引中未找到相關內容，請嘗試換不同詞語提問）"]
            else:
                groups: dict[str, list[Chunk]] = {}
                for c in top_chunks:
                    groups.setdefault(c.source, []).append(c)
                for fname, chunks in groups.items():
                    sections = [
                        f"{c.page_tag}\n{c.text}" if c.page_tag else c.text
                        for c in chunks
                    ]
                    context_parts.append(
                        f"【{fname}】\n" + "\n…\n".join(sections)
                    )
                    source_files.append(fname)
        else:
            context_parts = ["（索引尚未建立，請稍候或點擊「重新整理」）"]

        context = ("\n\n" + "─" * 40 + "\n\n").join(context_parts)

        system_msg = (
            "你是一位專業的學校事務助手，根據學校官方 DOCX 文件回答問題。\n"
            "規則：\n"
            "• 只根據所提供的文件內容回答，不可憑空猜測。\n"
            "• 使用繁體中文，語氣親切、專業。\n"
            "• 若文件中找不到答案，請如實說明，並建議聯絡學校查詢。\n"
            "• 回答需條理清晰，適當使用列表或分段。\n"
            "• 如有引用，請說明出自哪個文件及頁數。"
        )

        messages: list[dict] = [{"role": "system", "content": system_msg}]
        if chat_history:
            for msg in chat_history[-6:]:
                messages.append({"role": msg["role"], "content": msg["content"]})
        messages.append({
            "role":    "user",
            "content": f"問題：{question}\n\n相關文件內容：\n{context}",
        })

        try:
            resp = self.client.chat.completions.create(
                model=self.model, messages=messages, max_tokens=2000, temperature=0.7
            )
            return resp.choices[0].message.content, source_files
        except Exception as e:
            raise ChatbotError(f"AI 回答失敗：{e}")
